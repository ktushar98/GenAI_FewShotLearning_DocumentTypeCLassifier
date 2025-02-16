'''import streamlit as st
import fitz  # PyMuPDF for PDF extraction
import docx
import torch
import json
from sentence_transformers import SentenceTransformer, util

# Load a pre-trained sentence transformer model
model = SentenceTransformer("all-MiniLM-L6-v2")

# --- LOAD DOCUMENT EXAMPLES FROM JSON ---
with open("few_shot_examples.json", "r", encoding="utf-8") as file:
    document_examples = json.load(file)

# Encode example documents for similarity comparison
example_embeddings = {
    doc_type: model.encode(description, convert_to_tensor=True)
    for doc_type, description in document_examples.items()
}

# --- FUNCTIONS ---
def extract_text_from_pdf(file):
    """Extract text from a PDF file."""
    doc = fitz.open(stream=file.read(), filetype="pdf")  
    text = "\n".join([page.get_text("text") for page in doc])
    return text if text.strip() else "No text found in PDF."

def extract_text_from_docx(file):
    """Extract text from a Word document."""
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def classify_document(text):
    """Classify document using Few-Shot and Zero-Shot Learning."""
    input_embedding = model.encode(text, convert_to_tensor=True)
    
    # Calculate similarity scores
    similarities = {
        doc_type: util.pytorch_cos_sim(input_embedding, emb)[0].item()
        for doc_type, emb in example_embeddings.items()
    }

    predicted_category = max(similarities, key=similarities.get)
    confidence_score = round(similarities[predicted_category] * 100, 2)

    # Zero-Shot Handling
    if confidence_score < 40:
        explanation = (
            "This document does not strongly match any known category. "
            "Providing the closest potential classifications."
        )
        closest_matches = sorted(similarities.items(), key=lambda x: x[1], reverse=True)[:3]
        predicted_category = "Uncertain"
        suggestions = [f"{doc}: {round(score * 100, 2)}%" for doc, score in closest_matches]
        return predicted_category, confidence_score, explanation, suggestions
    else:
        explanation = (
            f"This document is classified as **{predicted_category}** based on content structure, "
            f"terminology, and format. The high confidence score suggests a strong match."
        )
        return predicted_category, confidence_score, None, None

# --- STREAMLIT APP ---
st.title("📄Business Document Classifier")
st.write("Upload a PDF or Word document to classify its type using Few-Shot & Zero-Shot Learning.")

uploaded_file = st.file_uploader("Upload your document (PDF/DOCX)", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("Extracting text..."):
        if uploaded_file.type == "application/pdf":
            extracted_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_text = extract_text_from_docx(uploaded_file)
        else:
            st.error("Unsupported file format.")
            st.stop()

    if extracted_text and extracted_text != "No text found in PDF.":
        with st.spinner("Classifying document..."):
            predicted_category, confidence_score, explanation, suggestions = classify_document(extracted_text)

        # --- OUTPUT RESULTS ---
        st.subheader("📌 Classification Results")
        st.write(f"**Document Category:** {predicted_category}")
        st.write(f"**Confidence Score:** {confidence_score}%")

        if explanation:
            st.write(f"**Explanation:** {explanation}")

        if suggestions:
            st.subheader("🔍 Possible Matches")
            for match in suggestions:
                st.write(f"- {match}")

    else:
        st.error("No readable text found in the document.")
'''

import streamlit as st
import fitz  # PyMuPDF for PDF extraction
import docx
import torch
import json
from sentence_transformers import SentenceTransformer, util

# Load a pre-trained sentence transformer model
model = SentenceTransformer("all-MiniLM-L6-v2")

# --- LOAD DOCUMENT EXAMPLES FROM JSON ---
with open("few_shot_examples.json", "r", encoding="utf-8") as file:
    document_examples = json.load(file)

# Encode example documents for similarity comparison
example_embeddings = {
    doc_type: model.encode(description, convert_to_tensor=True)
    for doc_type, description in document_examples.items()
}

# --- FUNCTIONS ---
def extract_text_from_pdf(file):
    """Extract text from a PDF file."""
    doc = fitz.open(stream=file.read(), filetype="pdf")  
    text = "\n".join([page.get_text("text") for page in doc])
    return text if text.strip() else "No text found in PDF."

def extract_text_from_docx(file):
    """Extract text from a Word document."""
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def classify_document(text):
    """Classify document using Few-Shot and Zero-Shot Learning."""
    input_embedding = model.encode(text, convert_to_tensor=True)
    
    # Calculate similarity scores
    similarities = {
        doc_type: util.pytorch_cos_sim(input_embedding, emb)[0].item()
        for doc_type, emb in example_embeddings.items()
    }

    predicted_category = max(similarities, key=similarities.get)
    confidence_score = round(similarities[predicted_category] * 100, 2)

    # Zero-Shot Handling
    if confidence_score < 40:
        explanation = (
            "This document does not strongly match any known category. "
            "Providing the closest potential classifications."
        )
        closest_matches = sorted(similarities.items(), key=lambda x: x[1], reverse=True)[:3]
        predicted_category = "Uncertain"
        suggestions = [f"{doc}: {round(score * 100, 2)}%" for doc, score in closest_matches]
        return predicted_category, confidence_score, explanation, suggestions
    else:
        explanation = (
            f"This document is classified as **{predicted_category}** based on content structure, "
            f"terminology, and format. The high confidence score suggests a strong match."
        )
        return predicted_category, confidence_score, None, None

# --- STREAMLIT APP UI ---
st.set_page_config(page_title="Business Document Classifier", page_icon="📄", layout="wide")

# Header
st.markdown(
    """
    <style>
        .title {
            font-size: 36px;
            font-weight: bold;
            text-align: center;
            color: #2c3e50;
        }
        .sub {
            font-size: 20px;
            text-align: center;
            color: #7f8c8d;
        }
    </style>
    """,
    unsafe_allow_html=True
)
st.markdown('<p class="title">📄 Business Document Classifier</p>', unsafe_allow_html=True)
st.markdown('<p class="sub">Upload a PDF or Word document to classify its type using Few-Shot & Zero-Shot Learning.</p>', unsafe_allow_html=True)

st.markdown("---")

# File upload section
uploaded_file = st.file_uploader("📂 Upload your document (PDF/DOCX)", type=["pdf", "docx"])

# Submit button
submit_button = st.button("🚀 Submit & Classify")

if submit_button:
    if uploaded_file is None:
        st.error("⚠️ Please upload a document before submitting.")
    else:
        with st.spinner("Extracting text..."):
            if uploaded_file.type == "application/pdf":
                extracted_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                extracted_text = extract_text_from_docx(uploaded_file)
            else:
                st.error("❌ Unsupported file format.")
                st.stop()

        if extracted_text and extracted_text != "No text found in PDF.":

            with st.spinner("🔍 Classifying document..."):
                predicted_category, confidence_score, explanation, suggestions = classify_document(extracted_text)

            # --- OUTPUT RESULTS ---
            st.markdown("---")
            st.subheader("📌 Classification Results")

            st.success(f"**✅ Document Category:** {predicted_category}")
            st.info(f"📊 **Confidence Score:** {confidence_score}%")

            if explanation:
                st.warning(f"💡 **Explanation:** {explanation}")

            if suggestions:
                st.subheader("🔍 Possible Matches")
                for match in suggestions:
                    st.write(f"- {match}")

            # Display extracted text in a collapsible section
            with st.expander("📜 View Extracted Document Text"):
                st.text_area("Extracted Text", extracted_text, height=300)

        else:
            st.error("⚠️ No readable text found in the document.")

# Footer
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: grey;'>"
    "🔹 Built with ❤️ using Streamlit | GenAI-powered Classification System 🔹"
    "</div>",
    unsafe_allow_html=True
)
